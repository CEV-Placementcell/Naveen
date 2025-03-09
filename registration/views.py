from django.shortcuts import render,redirect, get_object_or_404
from docx.shared import Inches
from docx import Document
import random, io
from .models import student, contact
from . models import *
from management.models import *
from datetime import date
from django.core.mail import send_mail,EmailMessage
from django.conf import settings
from django.http import  FileResponse,HttpResponse,HttpResponseRedirect
from django.views.decorators.csrf import csrf_exempt
from exam import models as QMODEL
from django.contrib import messages
#from .forms import StudentEditForm, StudentContactEditForm
# In views.py
from control.models import Card
from control.models import GalleryImage
from django.core.files.storage import FileSystemStorage
from control.models import DriveDetails


def testimonials_view(request):
    # Fetch all cards for display
    cards = Card.objects.all()
    return render(request, 'testimonials.html', {'cards': cards})

def gallery_view(request):
    gallery_images = GalleryImage.objects.all()
    return render(request, 'gallery.html', {'gallery_images': gallery_images})

def drive_view(request):
    """
    View to display placement statistics from the Drive model.
    """
    placements = Drive.objects.all().order_by('-year')  # Fetch all placement records sorted by year (latest first)
    return render(request, "drive.html", {"placements": placements})



def drive_details(request, year):
    drives = DriveDetails.objects.filter(year=year)
    return render(request, 'drivedetails.html', {'drives': drives, 'year': year})


# Create your views here.
def index(request):
    notifify=notification.objects.all()
    return render(request,'index.html',{"notifify":notifify})

def registration(request):
    if request.method=="POST":
        name=request.POST['name']
        ad_no= request.POST['admission_no']
        sslc= request.POST['sslc']
        dob= request.POST['dob']
        yo_add = request.POST['yo_add']
        dept= request.POST['dept']
        prog= request.POST['program']
        course= request.POST['course']
        phone_no= request.POST['phone_no']
        emailid=request.POST['email']
        address= request.POST['address']
        fname= request.POST['fname']
        mname= request.POST['mname']
        dis= request.POST['dis']
        state= request.POST['state']
        pin= request.POST['pin']
        g_ph= request.POST['g_number']
        a_ins= request.POST['a_ins']
        skill=request.POST['skill']
        photo=request.FILES['photo']
        password=request.POST['password']
        cpassword=request.POST['cpassword']
        aadhar=request.POST['aadhar']
        hsc=request.POST['hsc']
        gpa=request.POST['gpa']

        if password==cpassword:
            u = student(name=name,ad_no=ad_no,sslc=sslc,dob=dob,yo_add=yo_add,dept=dept,prog=prog,hsc=hsc,
            aadhar=aadhar,gpa=gpa,course=course,stud_ph=phone_no,area_int=a_ins,skill=skill,photo=photo,password=password)

            u.save()

            user = contact(ad_no_id=ad_no,email=emailid,adr=address,f_name=fname,m_name=mname,
            dist=dis,st=state,pin=pin,gua_ph=g_ph)
            user.save()

            subject= "Placement Cell: Registartion Successfull"
            msg = "Dear "+name+",\n\nWe appreciate Your interest in registering in Placement Cell of College of Engineering - Vadakara, Mandarathur.It is mandatory to submit the printout to complete the registration process.\n\n-Training and Placement Officer\nCollege of Engineering - Vadakara, Mandarathur."
            email_from= settings.EMAIL_HOST_USER
            reciver = [user.email]
            send_mail(subject,msg,email_from,reciver)
            send_email_view(request,ad_no)
            return redirect('/')

        else:
            messages.info(request,"Password mismatch, try again !!!")
            return redirect('registration')

    return render(request,'regold.html')


def studlogin(request):
    if request.method=="POST":
        try:
            ad_no= request.POST['admission_no']
            password= request.POST['password']
            # u = authenticate(request,ad_no=ad_no,password=password)
            # if u is not None:
            
            u=student.objects.get(ad_no=ad_no,password=password)
            #     login(request, u)
            # user = authenticate(request,ad_no=u.ad_no,password=u.password)
            # if user is not None:
            #     login(request, user)
            request.session['ad_no']=u.ad_no
            return redirect('studentinterface',ad_no)
        except student.DoesNotExist as e:
            messages.info(request,"Invalid Admission number or password")
            return redirect('studlogin')
    
    return render(request,'stlogin.html')

def studlogout(request):
    if 'ad_no' in request.session:
        request.session.flush()
    return redirect(studlogin)


def studentinterface(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)
        livedrives=job.objects.order_by("-d_no")

        return render(request,'studentinter.html',{"stud":stud,"livedrives":livedrives})
    return redirect(studlogin)

def posterview(request,id,dno):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        try:
            livedrives=job.objects.get(d_no=dno)
            posterpath=livedrives.poster.path
            return FileResponse(open(posterpath,'rb'))
        except:
            messages.info(request,"No Poster Uploaded")
            return redirect(studentinterface,id=id)



def applyjob(request,id,dno):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        try:
            appjob=jobs_applied.objects.get(d_no_id=dno,ad_no_id=id)
            messages.info(request,"Already applied")
            return(redirect(studentinterface,id=id))
            
        except jobs_applied.DoesNotExist as e :
            appjob=job.objects.get(d_no=dno)
            stud=student.objects.get(ad_no=id)

            d=date.today()
            newapplication=jobs_applied(d_no_id=dno,ad_no_id=id,date=d)
            newapplication.save()
            driveno=str(appjob.d_no)
            messages.info(request,"Thankyou for Submission")
            subject= "Placement Cell: Drive no: "+driveno+" Registration"
            msg = "Dear "+stud.name+",\n\nWe appreciate Your interest in registering for drive no: "+driveno+", for the position of  "+appjob.j_pos+" at "+appjob.c_name+".\n\n-Training and Placement Officer\nCollege of Engineering - Vadakara, Mandarathur."
            email_from= settings.EMAIL_HOST_USER
            reciver = [stud.contact.email]
            send_mail(subject,msg,email_from,reciver)

            return(redirect(studentinterface,id=id))


    return redirect(studlogin)


def stprofile(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)
        stud_cont=contact.objects.get(ad_no=id)
        yop=stud.yo_add+3


        return render(request,'profile.html',{"stud":stud,"st_con":stud_cont,"yop":yop})
    return redirect(studlogin)

def edit_profile(request, ad_no):
    # Check if 'ad_no' is in session and matches the requested `ad_no`
    if 'ad_no' in request.session and request.session['ad_no'] == ad_no:
        # Fetch user details
        user = get_object_or_404(student, ad_no=ad_no)
        contact_info = get_object_or_404(contact, ad_no_id=ad_no)

        if request.method == 'POST':
            # Update student details
            user.name = request.POST.get('name', user.name)
            user.sslc = request.POST.get('sslc', user.sslc)
            user.dob = request.POST.get('dob', user.dob)
            user.yo_add = request.POST.get('yo_add', user.yo_add)
            user.dept = request.POST.get('dept', user.dept)
            user.prog = request.POST.get('prog', user.prog)
            user.course = request.POST.get('course', user.course)
            user.stud_ph = request.POST.get('phone_no', user.stud_ph)
            user.area_int = request.POST.get('a_ins', user.area_int)
            user.skill = request.POST.get('skill', user.skill)
            user.aadhar = request.POST.get('aadhar', user.aadhar)
            user.hsc = request.POST.get('hsc', user.hsc)
            user.gpa = request.POST.get('gpa', user.gpa)

            # **Photo Update Logic**
            if 'photo' in request.FILES:  # If a new photo is uploaded
                if user.photo and user.photo.name != 'NULL':  
                    user.photo.delete(save=False)  # Delete the old file
                user.photo = request.FILES['photo']  # Assign new file

            # Password update section
            new_password = request.POST.get('password')
            confirm_password = request.POST.get('cpassword')

            if new_password and confirm_password:
                if new_password == confirm_password:
                    user.password = new_password  # Store the password directly without hashing
                else:
                    messages.error(request, "Passwords do not match!")
                    return render(request, 'edit_profile.html', {'user': user, 'contact_info': contact_info})

            user.save()

            # Update contact details
            contact_info.email = request.POST.get('email', contact_info.email)
            contact_info.adr = request.POST.get('address', contact_info.adr)
            contact_info.f_name = request.POST.get('fname', contact_info.f_name)
            contact_info.m_name = request.POST.get('mname', contact_info.m_name)
            contact_info.dist = request.POST.get('dis', contact_info.dist)
            contact_info.st = request.POST.get('state', contact_info.st)
            contact_info.pin = request.POST.get('pin', contact_info.pin)
            contact_info.gua_ph = request.POST.get('g_number', contact_info.gua_ph)

            contact_info.save()
            if 'photo' in request.FILES:
                print("✅ Photo received:", request.FILES['photo'])
            else:
                print("❌ No new photo uploaded")

            messages.success(request, "Profile updated successfully!")
            return redirect(f'/stprofile/{ad_no}/edit_profile')

    else:
        messages.error(request, "Unauthorized access!")
        return redirect('/studlogin')

    return render(request, 'edit_profile.html', {'user': user, 'contact_info': contact_info})



def apphistory(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)
        applyevent=events_applied.objects.filter(ad_no_id=id)
        events=event.objects.all()
        applyjob=jobs_applied.objects.filter(ad_no_id=id)
        jobs=job.objects.all()
        placed=placements.objects.filter(ad_no_id=id)

        return render(request,'apphistory.html',{"stud":stud,"applyevent":applyevent,"event":events,"applyjob":applyjob,"jobs":jobs,"placed":placed})
    return redirect(studlogin)

def apphistory2(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)
        applyevent=events_applied.objects.filter(ad_no_id=id)
        events=event.objects.all()
        applyjob=jobs_applied.objects.filter(ad_no_id=id)
        jobs=job.objects.all()
        placed=placements.objects.filter(ad_no_id=id)

        return render(request,'apphistory2.html',{"stud":stud,"applyevent":applyevent,"event":events,"applyjob":applyjob,"jobs":jobs,"placed":placed})
    return redirect(studlogin)

def apphistory3(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)
        applyevent=events_applied.objects.filter(ad_no_id=id)
        events=event.objects.all()
        applyjob=jobs_applied.objects.filter(ad_no_id=id)
        jobs=job.objects.all()
        placed=placements.objects.filter(ad_no_id=id)

        return render(request,'apphistory3.html',{"stud":stud,"applyevent":applyevent,"event":events,"applyjob":applyjob,"jobs":jobs,"placed":placed})
    return redirect(studlogin)
    


def eventreg(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)

        events = event.objects.all()
        try:
            ehistory = events_applied.objects.filter(ad_no=id)
        except:
            ehistory="NULL"

        return render(request,'eventreg.html',{"stud":stud,"events":events,"ehistory":ehistory})
    return redirect(studlogin)

def applyevent(request,id,eid):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        try:
            appevent=events_applied.objects.get(e_id_id=eid,ad_no_id=id)
            messages.info(request,"Already applied")
            return(redirect(eventreg,id=id))
            
        except events_applied.DoesNotExist as e :
            appevent=event.objects.get(e_id=eid)
            stud=student.objects.get(ad_no=id)
            d=date.today()
            newapplication=events_applied(e_id_id=eid,ad_no_id=id,date=d)
            newapplication.save()
            edate=str(appevent.date)
            messages.info(request,"Application Submitted")
            subject= "Placement Cell: "+appevent.e_name+" Registration"
            msg = "Dear "+stud.name+",\n\nWe appreciate Your interest in registering in "+appevent.e_name+", held on "+edate+" at "+appevent.venue+".\n\n-Training and Placement Officer\nCollege of Engineering - Vadakara, Mandarathur."
            email_from= settings.EMAIL_HOST_USER
            reciver = [stud.contact.email]
            send_mail(subject,msg,email_from,reciver)

            return(redirect(eventreg,id=id))


    return redirect(studlogin)


def doubt(request,id):
    if 'ad_no' in request.session and request.session['ad_no']==id :
        stud=student.objects.get(ad_no=id)
        try:
            if request.method=="POST":
                ad_no= request.POST['adno']
                d_no= request.POST['Driveno']
                d_title= request.POST['drive_title']
                d_descr= request.POST['drive_desc']
                d_ss= request.FILES['img']
                drive=job.objects.get(d_no=d_no)
                d= query(ad_no_id=ad_no,d_no_id=d_no,d_title=d_title,d_descr=d_descr,d_ss=d_ss)
                d.save()
        except job.DoesNotExist as e:
            messages.info(request,"Drive not found")
    

        return render(request,'studentsupport.html',{"stud":stud})
    return redirect(studlogin)

def studresponses(request,id):
    if 'ad_no' in request.session :
        stud=student.objects.get(ad_no=id)
        doubt=query.objects.filter(ad_no_id=id)
        return render(request,'supportresponses.html',{"doubt":doubt,"stud":stud})
    return redirect(studlogin)

def deletereplay(request,dno,ad_no,id):
   r=query.objects.get(d_no_id=dno,ad_no_id=ad_no,id=id)
   r.delete()
   return(redirect(studresponses,id=ad_no))


# def placed(request,id,dno):
#     if 'ad_no' in request.session and request.session['ad_no']==id :
#         stud=student.objects.get(ad_no=id)
#         jobs=job.objects.get(d_no=dno)
#         if request.method=="POST":
#             ad_no= request.POST['adno']
#             d_no= request.POST['dno']
#             p=placements(ad_no_id=ad_no,d_no_id=d_no)
#             p.save()
        

#         return render(request,'Placed.html',{"stud":stud,"jobs":jobs})
#     return redirect(studlogin)
    



def student_dashboard_view(request, id):
    if 'ad_no' in request.session:

        dict = {
            'stud': student.objects.get(ad_no=id),
            'total_course': QMODEL.Course.objects.all().count(),
            'total_question': QMODEL.Question.objects.all().count(),
        }
        return render(request, 'student_dashboard.html', context=dict)
    return redirect(studlogin)



def generate_pdf(id):
    document_path = 'static/registration-form.docx'
    document = Document(document_path)
    stud = student.objects.get(ad_no=id)

    # Construct the path to the photo file
    photo_path = stud.photo

    # Loop through each paragraph in the document
    for paragraph in document.paragraphs:

        paragraph.text = paragraph.text.replace('{stud.name}', stud.name)
        paragraph.text = paragraph.text.replace('{stud.dob}', str(stud.dob))
        paragraph.text = paragraph.text.replace('{stud.prog}', stud.prog)
        paragraph.text = paragraph.text.replace('{stud.course}', stud.course)
        paragraph.text = paragraph.text.replace('{stud.dept}', stud.dept)
        paragraph.text = paragraph.text.replace('{stud.stud_ph}', stud.stud_ph)
        paragraph.text = paragraph.text.replace('{stud.contact.email}', stud.contact.email)
        paragraph.text = paragraph.text.replace('{stud.ad_no}', stud.ad_no)


        # Find the placeholder text and replace it with an image
        if '{stud.photo}' in paragraph.text:
            # Remove the placeholder text
            paragraph.clear()
            paragraph.add_run().add_picture(photo_path, width=Inches(1.25))

    output = io.BytesIO()
    document.save(output)

    return output.getvalue()

@csrf_exempt
def send_email_view(request,id):
    try:
        stud = student.objects.get(ad_no=id)
        stud_cont = contact.objects.get(ad_no=id)
        subject = f'Placement Cell : Registration Form ({stud.ad_no})'
        message = f'Hi {stud.name}, you have been successfully registered in Placement Cell with Admission Number {stud.ad_no}.'
        email_from = settings.EMAIL_HOST_USER
        recipient_list = [stud_cont.email, ]
        email = EmailMessage(subject, message, email_from, recipient_list,)

        # get the image file and read it into memory
        if not stud.send:
            pdf_data = generate_pdf(id)
            email.attach('registration_report.docx',
                         pdf_data, 'application/docx')
            email.send()
            stud.send = True
            stud.save()
        else:
            messages.info(
                request, f"PDF already sent for Admission Number {stud.ad_no}")
    except Exception as e:
        messages.error(request, f"An error occurred while sending email: {e}")
        print(f'Error sending email: {str(e)}')
        stud.send = False
        stud.save()
    finally:
        return HttpResponse('Email sent')
    

def student_exam_view(request, id):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=id)
        courses = QMODEL.Course.objects.all()
        return render(request, 'student_exam.html', {'stud': stud, 'courses': courses})
    return redirect(studlogin)


def student_marks_view(request, id):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=id)
        courses = QMODEL.Course.objects.all()
        return render(request, 'student_marks.html', {'stud': stud, 'courses': courses})
    return redirect(studlogin)


def student_view_solution(request, id):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=id)
        key = QMODEL.Course.objects.all()
        return render(request, 'view_solution.html', {'stud': stud, 'key': key})
    return redirect(studlogin)


def take_exam_view(request, pk, ck):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=ck)
        course = QMODEL.Course.objects.get(id=pk)
        total_questions = QMODEL.Question.objects.all().filter(course=course).count()
        questions = QMODEL.Question.objects.all().filter(course=course)
        total_marks = 0
        for q in questions:
            total_marks = total_marks + q.marks
        return render(request, 'take_exam.html', {'stud': stud, 'course': course, 'total_questions': total_questions, 'total_marks': total_marks})
    return redirect(studlogin)


def check_marks_view(request, pk, ck):
    if 'ad_no' in request.session:
        course = QMODEL.Course.objects.get(id=pk)
        stud = student.objects.get(ad_no=ck)
        results = QMODEL.Result.objects.all().filter(exam=course).filter(student=stud)
        return render(request, 'check_marks.html', {'course':course,'stud':stud, 'results': results})
    return redirect(studlogin)


def download_solution(request, id, pk):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=pk)
        course = get_object_or_404(QMODEL.Course, id=id)
        file_path = course.solution.path
        return FileResponse(open(file_path, 'rb'))
    return redirect(studlogin)


def start_exam_view(request, pk, ck):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=ck)
        course = QMODEL.Course.objects.get(id=pk)
        qprint = course.qprint
        duration = course.time * 60
        questions = list(QMODEL.Question.objects.all().filter(course=course))
        questions = random.sample(questions, qprint)
        if request.method == 'POST':
            pass
        response = render(request, 'start_exam.html', {
                          'stud': stud, 'course': course, 'questions': questions, 'duration': duration})
        response.set_cookie('course_id', course.pk)
        return response
    return redirect(studlogin)


@csrf_exempt
def calculate_marks_view(request, pk):
    if 'ad_no' in request.session:
        if request.COOKIES.get('course_id') is not None:
            course_id = request.COOKIES.get('course_id')
            course = QMODEL.Course.objects.get(id=course_id)

            total_marks = 0
            questions = QMODEL.Question.objects.all().filter(course=course)
            for i in range(len(questions)):
                selected_ans = request.COOKIES.get(str(i+1))
                actual_answer = questions[i].answer
                if selected_ans == actual_answer:
                    total_marks = total_marks + questions[i].marks

            stud = student.objects.get(ad_no=pk)
            result = QMODEL.Result()
            result.marks = total_marks
            result.exam = course
            result.student = stud
            result.save()

            return HttpResponseRedirect('/view-result/{}'.format(pk))
    return redirect(studlogin)

def view_result_view(request,pk):
    if 'ad_no' in request.session:
        stud = student.objects.get(ad_no=pk)
        courses=QMODEL.Course.objects.all()
        return render(request,'view_result.html',{'stud':stud,'courses':courses})
    return redirect(studlogin)




